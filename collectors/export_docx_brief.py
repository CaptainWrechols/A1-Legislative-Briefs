#!/usr/bin/env python3
"""Build the front-brief .docx with direct formatting (python-docx).

Unlike the pandoc route (which formats via style definitions that Word,
Google Docs, and LibreOffice each resolve differently), this writer puts
literal formatting on every run: explicit Arial, explicit RGB colors, and
real uppercase text instead of the all-caps style property. The file then
looks the same in Microsoft Word, Word Online, Google Docs, LibreOffice,
and Apple Pages.

Layout mirrors the Phase 2 Issue Brief system and the packaged HTML:
masthead ("THE FORUM" + navy rule), navy title, terracotta ALL-CAPS section
headers, a Key-numbers stat strip (single-row card table with navy top
bars), justified 9pt body with navy bold lead-ins, and a gray small-caps
footline.

  python collectors/export_docx_brief.py --brief-dir briefs/nevada/water-scarcity/citizen-v1
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1A, 0x2D, 0x4F)
NAVY2 = RGBColor(0x2E, 0x4A, 0x78)
TERRACOTTA = RGBColor(0xC0, 0x39, 0x2B)
BODY = RGBColor(0x44, 0x44, 0x44)
MUTED = RGBColor(0x66, 0x66, 0x66)
CARD_BORDER = "CCCCCC"


def set_font(run, size: float, color: RGBColor, bold: bool = False, spacing_pts: float | None = None):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attr), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    if spacing_pts:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(int(spacing_pts * 20)))
        rpr.append(sp)


def para(doc, before=0.0, after=4.0, align=None, line=1.06):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.widow_control = True
    if align is not None:
        pf.alignment = align
    return p


def set_columns(section, num: int, space_twips: int = 240) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def no_widow(text: str) -> str:
    """Glue the final words together with non-breaking spaces so the last
    line of a paragraph never holds fewer than three words (the glued run
    is kept short so justification gaps stay invisible)."""
    m = re.search(r"(\S+) (\S+) (\S+)$", text)
    if m and sum(len(g) for g in m.groups()) <= 22:
        return text[: m.start()] + "\u00a0".join(m.groups())
    m = re.search(r"(\S+) (\S+)$", text)
    if m and len(m.group(1)) + len(m.group(2)) <= 24:
        return text[: m.start()] + m.group(1) + "\u00a0" + m.group(2)
    return text


def add_bottom_border(p, color: str, size_eighth_pts: int):
    ppr = p._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighth_pts))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    ppr.append(borders)


def add_top_border(p, color: str, size_eighth_pts: int):
    ppr = p._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size_eighth_pts))
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), color)
    borders.append(top)
    ppr.append(borders)


def set_cell_borders(cell, top_color: str, top_size: int, side_color: str):
    tcpr = cell._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, color, size in (
        ("top", top_color, top_size),
        ("left", side_color, 4),
        ("bottom", side_color, 4),
        ("right", side_color, 4),
    ):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcpr.append(borders)
    margins = OxmlElement("w:tcMar")
    for edge, tw in (("top", 50), ("left", 90), ("bottom", 60), ("right", 90)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(tw))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tcpr.append(margins)


def add_rich_text(p, text: str, size=9.0, color=BODY, bold_color=NAVY):
    """Write text with **bold** segments rendered bold in bold_color."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, size, bold_color, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, size, color)


def _is_explainer(line: str) -> bool:
    """Full-line single-asterisk italics used as section explainers."""
    s = line.strip()
    return (
        len(s) > 2
        and s.startswith("*")
        and s.endswith("*")
        and not s.startswith("**")
        and not s.endswith("**")
    )


def parse_markdown(md_path: Path):
    text = md_path.read_text(encoding="utf-8")
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            text = text[end + 4 :]
    text = re.sub(r"<!--.*?-->\n?", "", text, flags=re.S)
    lines = [l.rstrip() for l in text.splitlines()]

    title = ""
    subtitle = ""
    sections: list[dict] = []
    footline = ""
    current = None
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# ") and not title:
            title = line[2:].strip()
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and not lines[j].startswith("#"):
                subtitle = lines[j].strip()
                i = j
        elif line.startswith("### ") and current is not None:
            current["items"].append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            current = {"heading": line[3:].strip(), "items": []}
            sections.append(current)
        elif line.strip() == "---":
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            footline = " ".join(l.strip() for l in lines[j:] if l.strip())
            break
        elif line.strip().startswith("- ") and current is not None:
            current["items"].append(("bullet", line.strip()[2:].strip()))
        elif _is_explainer(line) and current is not None:
            current["items"].append(("explainer", line.strip()[1:-1].strip()))
        elif line.strip() and current is not None:
            block = [line.strip()]
            while (
                i + 1 < len(lines)
                and lines[i + 1].strip()
                and not lines[i + 1].startswith(("#", "- "))
                and lines[i + 1].strip() != "---"
                and not _is_explainer(lines[i + 1])
            ):
                i += 1
                block.append(lines[i].strip())
            current["items"].append(("para", " ".join(block)))
        i += 1
    return title, subtitle, sections, footline


def build_stat_strip(doc, bullets: list[str]):
    """Bullets like '**108** caption text' become stat cards."""
    stats = []
    for b in bullets:
        m = re.match(r"\*\*(.+?)\*\*\s*(.*)", b)
        if m:
            caption = m.group(2).lstrip("—-– ").strip()
            stats.append((m.group(1), caption))
    if not stats:
        return
    table = doc.add_table(rows=1, cols=len(stats))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, (num, caption) in enumerate(stats):
        cell = table.cell(0, idx)
        set_cell_borders(cell, top_color="1A2D4F", top_size=24, side_color=CARD_BORDER)
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(1)
        run = p1.add_run(num)
        set_font(run, 18 if len(num) <= 4 else 14, NAVY, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        run = p2.add_run(caption)
        set_font(run, 10, MUTED)
    spacer = para(doc, after=2)
    spacer.paragraph_format.space_after = Pt(2)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--brief-dir", required=True)
    parser.add_argument("--file", default="citizen-brief",
                        help="basename of the markdown/docx pair (default citizen-brief)")
    parser.add_argument("--layout", choices=["brief", "glossary", "prose"], default=None,
                        help="brief: justified 2-page front-brief; glossary: two-column "
                             "Term: entries; prose: single-column left-aligned companion "
                             "at 1.15 line spacing")
    args = parser.parse_args()
    layout = args.layout or ("brief" if args.file == "citizen-brief" else "glossary")
    brief_dir = Path(args.brief_dir)
    md_path = brief_dir / f"{args.file}.md"
    out_path = brief_dir / f"{args.file}.docx"

    keep_entries_together = layout == "glossary"
    prose = layout == "prose"

    title, subtitle, sections, footline = parse_markdown(md_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(0.6))

    # Masthead: THE FORUM + navy rule
    p = para(doc, after=1)
    run = p.add_run("T H E   F O R U M")
    set_font(run, 10, NAVY, bold=True, spacing_pts=1.0)
    rule = para(doc, after=6)
    add_bottom_border(rule, "1A2D4F", 20)

    # Title + subtitle
    p = para(doc, after=2, line=1.05)
    run = p.add_run(title)
    set_font(run, 18, NAVY, bold=True)
    if subtitle:
        p = para(doc, after=6)
        add_rich_text(p, no_widow(subtitle), size=10, color=BODY)


    for sec_index, sec in enumerate(sections):
        if keep_entries_together and sec_index > 0:
            head_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            set_columns(head_section, 1)
        p = para(doc, before=4.2, after=1.8)
        run = p.add_run(sec["heading"].upper())
        set_font(run, 12.5, TERRACOTTA, bold=True, spacing_pts=1.0)
        if keep_entries_together:
            entry_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            set_columns(entry_section, 2)
        stat_bullets = [
            t for kind, t in sec["items"]
            if kind == "bullet"
            and (m := re.match(r"\*\*(.+?)\*\*\s", t))
            and len(m.group(1)) <= 12
            and not m.group(1).endswith(":")
        ]
        for kind, text in sec["items"]:
            if kind == "bullet" and text in stat_bullets:
                continue
            if kind == "h3":
                p = para(doc, before=3.5, after=1.5, line=1.04)
                run = p.add_run(text)
                set_font(run, 12, NAVY2, bold=True)
            elif kind == "explainer":
                p = para(doc, after=3 if prose else 1.8, line=1.15 if prose else 1.03)
                add_rich_text(p, no_widow(text), size=10, color=MUTED)
                for run in p.runs:
                    run.font.italic = True
            elif kind == "bullet":
                p = para(doc, after=2.2 if keep_entries_together else 3,
                         line=1.0 if keep_entries_together else 1.05)
                p.paragraph_format.left_indent = Pt(10)
                run = p.add_run("▪  ")
                set_font(run, 10, TERRACOTTA)
                add_rich_text(p, no_widow(text), size=10)
            else:
                if prose:
                    p = para(doc, after=4.5, line=1.15)
                    p.paragraph_format.keep_together = True
                else:
                    p = para(doc, after=2.45, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.03)
                    if keep_entries_together:
                        p.paragraph_format.keep_together = True
                add_rich_text(p, no_widow(text), size=10)
        if stat_bullets:
            build_stat_strip(doc, stat_bullets)

    if footline:
        if keep_entries_together:
            foot_section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            set_columns(foot_section, 1)
        p = para(doc, before=2, after=0)
        add_top_border(p, "C8CDD6", 6)
        run = p.add_run(footline)
        set_font(run, 10, MUTED, spacing_pts=0.2)

    doc.save(out_path)
    print(f"Wrote {out_path} (direct-formatted, cross-app safe)")


if __name__ == "__main__":
    main()
